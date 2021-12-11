import fitz
import docx
import re

def read_pdf(file):
    text = ''
    with fitz.open(file) as doc:
        for page in doc:
            text+= re.sub('\s', '', page.get_text())
    return text



def read_docx(file):
    """
    读取docx文件，并返回其中的文本内容
    :param file: docx文件路径
    :return: docx中的文本内容
    """
    text = ''
    doc = docx.Document(file)
    for para in doc.paragraphs:
        text += para.text
    return text


"""
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
import re

def read_pdf(file):
    output_string = StringIO()
    with open(file, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
    text = output_string.getvalue()
    return text
"""
